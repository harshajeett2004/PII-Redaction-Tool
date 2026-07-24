from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from detector import PIIDetector
from fake_generator import FakeDataGenerator

import csv
import os


class PIIRedactor:

    # -------------------------------------------------------
    # Initialization
    # -------------------------------------------------------

    def __init__(self):

        self.detector = PIIDetector()
        self.fake_generator = FakeDataGenerator()

        self.logs = []

        self.stats = {
            "paragraphs": 0,
            "tables": 0,
            "headers": 0,
            "footers": 0,
            "entities": 0
        }

    # -------------------------------------------------------
    # Logging
    # -------------------------------------------------------

    def log_entity(
        self,
        original,
        fake,
        entity_type
    ):

        self.logs.append(
            {
                "Original": original,
                "Fake": fake,
                "Entity": entity_type
            }
        )

        self.stats["entities"] += 1

    # -------------------------------------------------------
    # Save CSV
    # -------------------------------------------------------

    def save_log(self, output_csv):

        folder = os.path.dirname(output_csv)

        if folder:
            os.makedirs(folder, exist_ok=True)

        with open(
            output_csv,
            "w",
            newline="",
            encoding="utf-8"
        ) as file:

            writer = csv.DictWriter(
                file,
                fieldnames=[
                    "Original",
                    "Fake",
                    "Entity"
                ]
            )

            writer.writeheader()

            writer.writerows(self.logs)

    # -------------------------------------------------------
    # Fake Value
    # -------------------------------------------------------

    def fake_value(self, entity):

        return self.fake_generator.generate(
            entity["type"],
            entity["text"]
        )

    # -------------------------------------------------------
    # Collect Every Paragraph
    # -------------------------------------------------------

    def collect_paragraphs(self, document):

        """
        Collect every paragraph from
        document,
        tables,
        headers,
        footers.

        This allows ONE batch NLP pass.
        """

        paragraphs = []

        # -----------------------------
        # Main document
        # -----------------------------

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                paragraphs.append(paragraph)

        # -----------------------------
        # Tables
        # -----------------------------

        for table in document.tables:

            self.stats["tables"] += 1

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        if paragraph.text.strip():

                            paragraphs.append(paragraph)

        # -----------------------------
        # Headers & Footers
        # -----------------------------

        for section in document.sections:

            headers = [
                section.header,
                section.first_page_header,
                section.even_page_header
            ]

            for header in headers:

                if header is None:
                    continue

                self.stats["headers"] += 1

                for paragraph in header.paragraphs:

                    if paragraph.text.strip():

                        paragraphs.append(paragraph)

                for table in header.tables:

                    for row in table.rows:

                        for cell in row.cells:

                            for paragraph in cell.paragraphs:

                                if paragraph.text.strip():

                                    paragraphs.append(paragraph)

            footers = [
                section.footer,
                section.first_page_footer,
                section.even_page_footer
            ]

            for footer in footers:

                if footer is None:
                    continue

                self.stats["footers"] += 1

                for paragraph in footer.paragraphs:

                    if paragraph.text.strip():

                        paragraphs.append(paragraph)

                for table in footer.tables:

                    for row in table.rows:

                        for cell in row.cells:

                            for paragraph in cell.paragraphs:

                                if paragraph.text.strip():

                                    paragraphs.append(paragraph)

        self.stats["paragraphs"] = len(paragraphs)

        return paragraphs
        # -------------------------------------------------------
    # Replace Inside Single Run
    # -------------------------------------------------------

    def replace_inside_run(
        self,
        run,
        original,
        replacement
    ):

        if original not in run.text:
            return False

        run.text = run.text.replace(
            original,
            replacement
        )

        return True

    # -------------------------------------------------------
    # Replace Across Multiple Runs
    # -------------------------------------------------------

    def replace_across_runs(
        self,
        paragraph,
        original,
        replacement
    ):

        full_text = ""

        run_positions = []

        current = 0

        for run in paragraph.runs:

            start = current
            end = start + len(run.text)

            run_positions.append(
                (
                    run,
                    start,
                    end
                )
            )

            full_text += run.text

            current = end

        start_index = full_text.find(original)

        if start_index == -1:
            return False

        end_index = start_index + len(original)

        affected = []

        for run, rs, re in run_positions:

            if rs < end_index and re > start_index:

                affected.append(
                    (
                        run,
                        rs,
                        re
                    )
                )

        if not affected:
            return False

        first_run = affected[0][0]
        first_start = affected[0][1]

        prefix = first_run.text[
            :start_index - first_start
        ]

        last_run = affected[-1][0]
        last_start = affected[-1][1]

        suffix = last_run.text[
            end_index - last_start:
        ]

        first_run.text = (
            prefix
            + replacement
            + suffix
        )

        for run, _, _ in affected[1:]:

            run.text = ""

        return True

    # -------------------------------------------------------
    # Replace Entity
    # -------------------------------------------------------

    def replace_entity(
        self,
        paragraph,
        entity
    ):

        fake = self.fake_value(entity)

        self.log_entity(
            entity["text"],
            fake,
            entity["type"]
        )

        for run in paragraph.runs:

            if self.replace_inside_run(
                run,
                entity["text"],
                fake
            ):
                return

        self.replace_across_runs(
            paragraph,
            entity["text"],
            fake
        )

    # -------------------------------------------------------
    # Batch Detection
    # -------------------------------------------------------

    def detect_all_entities(
        self,
        paragraphs
    ):

        print()

        print("=" * 60)
        print("Batch Detecting PII")
        print("=" * 60)

        texts = [
            paragraph.text
            for paragraph in paragraphs
        ]

        entity_lists = self.detector.detect_batch(
            texts
        )

        results = []

        for paragraph, entities in zip(
            paragraphs,
            entity_lists
        ):

            entities = sorted(

                entities,

                key=lambda x: len(
                    x["text"]
                ),

                reverse=True

            )

            results.append(
                (
                    paragraph,
                    entities
                )
            )

        print(
            f"Detection Completed "
            f"({len(texts)} paragraphs)"
        )

        return results

    # -------------------------------------------------------
    # Apply Redaction
    # -------------------------------------------------------

    def apply_redaction(
        self,
        paragraph,
        entities
    ):

        if not entities:
            return

        already_done = set()

        for entity in entities:

            key = (

                entity["type"],

                entity["text"]

            )

            if key in already_done:
                continue

            already_done.add(key)

            self.replace_entity(
                paragraph,
                entity
            )
    # -------------------------------------------------------
    # Batch Process Whole Document
    # -------------------------------------------------------

    def process_document(
        self,
        document
    ):

        print()
        print("=" * 70)
        print("COLLECTING DOCUMENT CONTENT")
        print("=" * 70)

        paragraphs = self.collect_paragraphs(
            document
        )

        print(
            f"Collected {len(paragraphs)} paragraphs."
        )

        print()
        print("=" * 70)
        print("RUNNING BATCH PII DETECTION")
        print("=" * 70)

        detected = self.detect_all_entities(
            paragraphs
        )

        print()
        print("=" * 70)
        print("APPLYING REDACTION")
        print("=" * 70)

        total = len(detected)

        for index, (
            paragraph,
            entities
        ) in enumerate(
            detected,
            start=1
        ):

            self.apply_redaction(
                paragraph,
                entities
            )

            if index % 100 == 0:

                print(
                    f"Redacted "
                    f"{index}/{total}"
                )

        print()

        print("=" * 70)
        print("DOCUMENT REDACTION COMPLETE")
        print("=" * 70)

    # -------------------------------------------------------
    # Print Summary
    # -------------------------------------------------------

    def print_summary(self):

        print()

        print("=" * 60)
        print("REDACTION SUMMARY")
        print("=" * 60)

        print(
            f"Paragraphs : "
            f"{self.stats['paragraphs']}"
        )

        print(
            f"Tables     : "
            f"{self.stats['tables']}"
        )

        print(
            f"Headers    : "
            f"{self.stats['headers']}"
        )

        print(
            f"Footers    : "
            f"{self.stats['footers']}"
        )

        print(
            f"Entities   : "
            f"{self.stats['entities']}"
        )

        print("=" * 60)

    # -------------------------------------------------------
    # Save Document
    # -------------------------------------------------------

    def save_document(
        self,
        document,
        output_path
    ):

        folder = os.path.dirname(
            output_path
        )

        if folder:

            os.makedirs(
                folder,
                exist_ok=True
            )

        document.save(output_path)

        print()

        print(
            "Saved Redacted Document:"
        )

        print(output_path)
        # -------------------------------------------------------
    # Main Redaction Pipeline
    # -------------------------------------------------------

    def redact_document(
        self,
        input_file,
        output_file,
        log_file="logs/redaction_log.csv"
    ):

        import time

        start_time = time.time()

        print()
        print("=" * 70)
        print("PII REDACTION STARTED")
        print("=" * 70)

        if not os.path.exists(input_file):
            raise FileNotFoundError(
                f"Input file not found:\n{input_file}"
            )

        print("\nLoading document...")

        document = Document(input_file)

        print("Document Loaded Successfully.")

        print()

        self.process_document(document)

        print()

        print("=" * 70)
        print("SAVING OUTPUT")
        print("=" * 70)

        self.save_document(
            document,
            output_file
        )

        self.save_log(
            log_file
        )

        self.print_summary()

        elapsed = time.time() - start_time

        minutes = int(elapsed // 60)
        seconds = elapsed % 60

        print()

        print("=" * 70)
        print("REDACTION COMPLETED")
        print("=" * 70)

        print(f"Output File : {output_file}")
        print(f"CSV Log     : {log_file}")

        print(
            f"Execution Time : "
            f"{minutes} min "
            f"{seconds:.2f} sec"
        )

        print("=" * 70)

        return output_file

    # -------------------------------------------------------
    # Get Statistics
    # -------------------------------------------------------

    def get_statistics(self):

        return {

            "paragraphs":
                self.stats["paragraphs"],

            "tables":
                self.stats["tables"],

            "headers":
                self.stats["headers"],

            "footers":
                self.stats["footers"],

            "entities":
                self.stats["entities"],

            "logs":
                len(self.logs)

        }

    # -------------------------------------------------------
    # Reset State
    # -------------------------------------------------------

    def reset(self):

        self.logs.clear()

        self.stats = {

            "paragraphs": 0,
            "tables": 0,
            "headers": 0,
            "footers": 0,
            "entities": 0

        }
        # -------------------------------------------------------
    # Export Logs as List
    # -------------------------------------------------------

    def get_logs(self):

        """
        Return all replacement logs.

        Useful for Streamlit dashboard.
        """

        return self.logs

    # -------------------------------------------------------
    # Total Entity Count
    # -------------------------------------------------------

    def total_entities(self):

        return len(self.logs)

    # -------------------------------------------------------
    # Check Empty Document
    # -------------------------------------------------------

    def is_empty_document(
        self,
        document
    ):

        paragraphs = self.collect_paragraphs(
            document
        )

        return len(paragraphs) == 0

    # -------------------------------------------------------
    # Validate Input File
    # -------------------------------------------------------

    def validate_input(
        self,
        input_file
    ):

        if not os.path.exists(input_file):

            raise FileNotFoundError(
                f"File not found:\n{input_file}"
            )

        extension = os.path.splitext(
            input_file
        )[1].lower()

        if extension != ".docx":

            raise ValueError(
                "Only DOCX files are supported."
            )

        return True


# ===========================================================
# MAIN
# ===========================================================

if __name__ == "__main__":

    INPUT_FILE = "input/Red Herring Prospectus.docx"

    OUTPUT_FILE = "output/Redacted_Prospectus.docx"

    LOG_FILE = "logs/redaction_log.csv"

    redactor = PIIRedactor()

    try:

        redactor.validate_input(
            INPUT_FILE
        )

        redactor.redact_document(

            INPUT_FILE,

            OUTPUT_FILE,

            LOG_FILE

        )

        print()

        print("=" * 70)

        print("PROCESS FINISHED SUCCESSFULLY")

        print("=" * 70)

        stats = redactor.get_statistics()

        print()

        print("Statistics")

        print("-" * 40)

        for key, value in stats.items():

            print(f"{key:<15}: {value}")

        print("-" * 40)

    except Exception as e:

        print()

        print("=" * 70)

        print("ERROR")

        print("=" * 70)

        print(str(e))

        print("=" * 70)
        # -------------------------------------------------------
    # Cached Fake Values
    # -------------------------------------------------------

    def get_fake_cached(self, entity):

        """
        Generate a fake value only once for each
        (entity type, original text) pair.
        """

        if not hasattr(self, "_fake_cache"):
            self._fake_cache = {}

        key = (
            entity["type"],
            entity["text"]
        )

        if key not in self._fake_cache:

            self._fake_cache[key] = self.fake_generator.generate(
                entity["type"],
                entity["text"]
            )

        return self._fake_cache[key]

    # -------------------------------------------------------
    # Replace Entity (Cached)
    # -------------------------------------------------------

    def replace_entity(
        self,
        paragraph,
        entity
    ):

        fake = self.get_fake_cached(entity)

        self.log_entity(
            entity["text"],
            fake,
            entity["type"]
        )

        for run in paragraph.runs:

            if self.replace_inside_run(
                run,
                entity["text"],
                fake
            ):
                return

        self.replace_across_runs(
            paragraph,
            entity["text"],
            fake
        )

    # -------------------------------------------------------
    # Clear Internal Cache
    # -------------------------------------------------------

    def clear_cache(self):

        if hasattr(self, "_fake_cache"):
            self._fake_cache.clear()